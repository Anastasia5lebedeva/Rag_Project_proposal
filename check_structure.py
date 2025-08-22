import os, sys, re
from pathlib import Path
from docx import Document

HEADING_RX = re.compile(r"^(heading|заголовок)\s*(\d+)?", re.IGNORECASE)

def heading_level(par):
    name = getattr(getattr(par, "style", None), "name", "") or ""
    m = HEADING_RX.match(name)
    if m:
        try: return int(m.group(2) or 1)
        except: return 1
    return 0

def preview_docx(path, max_paras=5, max_rows=3):
    p = Path(path).expanduser()
    if not p.exists():
        print(f"[ERR] file not found: {p}", file=sys.stderr); return 1
    try:
        doc = Document(str(p))
    except Exception as e:
        print(f"[ERR] cannot open DOCX: {e}", file=sys.stderr); return 2

    paras = list(doc.paragraphs)
    tables = list(doc.tables)

    headings = sum(1 for x in paras if heading_level(x) > 0)
    print(f"paras: {len(paras)} tables: {len(tables)}")
    print(f"headings by style: {headings}")

    print("\n-- first paragraphs --")
    for x in paras[:max_paras]:
        name = getattr(getattr(x, 'style', None), 'name', '') or ''
        print("-", x.text.strip(), "| style:", name)

    if tables:
        print("\n-- first rows from each table --")
        for ti, t in enumerate(tables, 1):
            print(f"[table {ti}]")
            for r in t.rows[:max_rows]:
                print("|".join(c.text.strip() for c in r.cells))

    return 0

if __name__ == "__main__":
    path = sys.argv[1] if len(sys.argv) > 1 else os.getenv("DOCX_PATH",
        "/home/anastasia/PycharmProjects/proposal-rag/MiningDetect/КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ.docx")
    sys.exit(preview_docx(path))